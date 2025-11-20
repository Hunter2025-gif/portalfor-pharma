from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.db.models import Q, Count
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from datetime import datetime, timedelta
from bmr.models import BMR, BMRSignature
from workflow.models import BatchPhaseExecution
import csv
import json

def get_filtered_comments_data(request):
    """Helper function to get comments data with filters applied"""
    # Check if user is admin/staff - they see all comments
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    # Collect all comments
    comments_data = []
    
    # 1. BMR Level Comments
    if is_admin:
        # Admin sees all BMRs
        bmrs = BMR.objects.all().select_related('product', 'created_by', 'approved_by')
    else:
        # Operators only see BMRs they created or were involved in
        bmrs = BMR.objects.filter(
            Q(created_by=request.user) | 
            Q(approved_by=request.user)
        ).select_related('product', 'created_by', 'approved_by')
    
    for bmr in bmrs:
        # QA Comments on BMR
        if bmr.qa_comments:
            comments_data.append({
                'bmr_number': bmr.batch_number,
                'product': bmr.product.product_name,
                'comment_type': 'BMR QA Comments',
                'phase': 'BMR Creation',
                'user': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
                'user_role': bmr.created_by.role if bmr.created_by else 'Unknown',
                'date': bmr.created_date,
                'comments': bmr.qa_comments,
                'status': bmr.status,
                'bmr_id': bmr.id
            })
        
        # Regulatory Comments on BMR
        if bmr.regulatory_comments:
            comments_data.append({
                'bmr_number': bmr.batch_number,
                'product': bmr.product.product_name,
                'comment_type': 'BMR Regulatory Comments',
                'phase': 'Regulatory Approval',
                'user': bmr.approved_by.get_full_name() if bmr.approved_by else 'Unknown',
                'user_role': bmr.approved_by.role if bmr.approved_by else 'Unknown',
                'date': bmr.approved_date or bmr.modified_date,
                'comments': bmr.regulatory_comments,
                'status': bmr.status,
                'bmr_id': bmr.id
            })
    
    # 2. Phase Level Comments
    if is_admin:
        # Admin sees all phases
        phases = BatchPhaseExecution.objects.all().select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    else:
        # Operators only see phases they were involved in
        phases = BatchPhaseExecution.objects.filter(
            Q(started_by=request.user) | 
            Q(completed_by=request.user) |
            Q(bmr__created_by=request.user)
        ).select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    
    for phase in phases:
        # Operator Comments
        if phase.operator_comments:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Operator Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'user_role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'comments': phase.operator_comments,
                'status': phase.status,
                'bmr_id': phase.bmr.id,
                'phase_id': phase.id
            })
        
        # QA Comments
        if phase.qa_comments:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Phase QA Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'user_role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'comments': phase.qa_comments,
                'status': phase.status,
                'bmr_id': phase.bmr.id,
                'phase_id': phase.id
            })
        
        # Rejection Reasons
        if phase.rejection_reason:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Rejection Reason',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'user_role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'comments': phase.rejection_reason,
                'status': phase.status,
                'bmr_id': phase.bmr.id,
                'phase_id': phase.id
            })
    
    # 3. Signature Comments
    if is_admin:
        # Admin sees all signatures
        signatures = BMRSignature.objects.all().select_related('bmr', 'bmr__product', 'signed_by')
    else:
        # Operators only see signatures they made or on BMRs they created
        signatures = BMRSignature.objects.filter(
            Q(signed_by=request.user) |
            Q(bmr__created_by=request.user)
        ).select_related('bmr', 'bmr__product', 'signed_by')
    
    for signature in signatures:
        if signature.comments:
            # Get user role from user's groups
            user_role = signature.signed_by.groups.first().name if signature.signed_by.groups.exists() else 'Staff'
            
            comments_data.append({
                'bmr_number': signature.bmr.batch_number,
                'product': signature.bmr.product.product_name,
                'comment_type': 'Electronic Signature',
                'phase': f"Signature - {signature.get_signature_type_display()}",
                'user': signature.signed_by.get_full_name() if signature.signed_by else 'Unknown',
                'user_role': user_role,
                'date': signature.signed_date,
                'comments': signature.comments,
                'status': 'Signed',
                'bmr_id': signature.bmr.id,
                'signature_id': signature.id
            })
    
    # Sort by date (newest first)
    comments_data.sort(key=lambda x: x['date'] or datetime.min, reverse=True)
    
    # Apply filters from the request
    bmr_filter = request.GET.get('bmr')
    comment_type_filter = request.GET.get('type')
    user_role_filter = request.GET.get('role')
    
    if bmr_filter:
        comments_data = [c for c in comments_data if bmr_filter.lower() in c['bmr_number'].lower()]
    
    if comment_type_filter:
        comments_data = [c for c in comments_data if c['comment_type'] == comment_type_filter]
    
    if user_role_filter:
        comments_data = [c for c in comments_data if c['user_role'] == user_role_filter]
    
    return comments_data

@login_required
def comments_report_view(request):
    """Web-based comments report view with role-based filtering and pagination"""
    
    from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
    
    # Check if user is admin/staff - they see all comments
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    # Collect all comments - either use the helper function or collect directly
    comments_data = get_filtered_comments_data(request)
    
    # Generate statistics
    stats = {
        'total_comments': len(comments_data),
        'bmr_comments': len([c for c in comments_data if 'BMR' in c['comment_type']]),
        'phase_comments': len([c for c in comments_data if c['comment_type'] in ['Operator Comments', 'Phase QA Comments']]),
        'rejections': len([c for c in comments_data if c['comment_type'] == 'Rejection Reason']),
        'signatures': len([c for c in comments_data if c['comment_type'] == 'Electronic Signature']),
    }
    
    # Get unique values for filters
    all_comment_types = list(set([c['comment_type'] for c in comments_data]))
    all_user_roles = list(set([c['user_role'] for c in comments_data]))
    all_bmrs = list(set([c['bmr_number'] for c in comments_data]))
    
    # Get current filters
    bmr_filter = request.GET.get('bmr')
    comment_type_filter = request.GET.get('type')
    user_role_filter = request.GET.get('role')
    
    # Pagination
    items_per_page = 15  # Set a reasonable number of items per page
    paginator = Paginator(comments_data, items_per_page)
    page = request.GET.get('page')
    
    try:
        paginated_comments = paginator.page(page)
    except PageNotAnInteger:
        # If page is not an integer, deliver first page
        paginated_comments = paginator.page(1)
    except EmptyPage:
        # If page is out of range, deliver last page
        paginated_comments = paginator.page(paginator.num_pages)
    
    context = {
        'comments': paginated_comments,
        'total_comments': len(comments_data),
        'stats': stats,
        'comment_types': sorted(all_comment_types),
        'user_roles': sorted(all_user_roles),
        'bmrs': sorted(all_bmrs),
        'current_filters': {
            'bmr': bmr_filter,
            'type': comment_type_filter,
            'role': user_role_filter
        },
        'is_admin': is_admin,  # Pass admin status to template
        'user_role': request.user.role  # Pass user role to template
    }
    
    return render(request, 'reports/comments_report.html', context)

@login_required
def export_comments_csv(request):
    """Export comments to CSV format with role-based filtering"""
    
    # Get filtered comments data
    raw_comments = get_filtered_comments_data(request)
    
    # Format data for CSV
    comments_data = []
    for comment in raw_comments:
        comments_data.append({
            'BMR Number': comment['bmr_number'],
            'Product': comment['product'],
            'Comment Type': comment['comment_type'],
            'Phase': comment['phase'],
            'Date': comment['date'].strftime('%Y-%m-%d %H:%M:%S') if comment['date'] else '',
            'Comments': comment['comments'],
            'Status': comment['status']
        })
    
    # Create CSV response
    response = HttpResponse(content_type='text/csv')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="KPI_Comments_Report_{timestamp}.csv"'
    
    if comments_data:
        writer = csv.DictWriter(response, fieldnames=comments_data[0].keys())
        writer.writeheader()
        writer.writerows(comments_data)
    else:
        writer = csv.writer(response)
        writer.writerow(['No comments found in the system'])
    
    return response

@login_required
def export_comments_word(request):
    """Export comments to Word format with role-based filtering"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        return HttpResponse("python-docx library not installed. Please install it to use Word export.", 
                          content_type="text/plain")
    
    try:
        # Get filtered comments data
        raw_comments = get_filtered_comments_data(request)

        # Group comments by BMR number
        bmr_comments = {}
        for comment in raw_comments:
            bmr_number = comment['bmr_number']
            if bmr_number not in bmr_comments:
                bmr_comments[bmr_number] = {
                    'product': comment['product'],
                    'comments': []
                }
            bmr_comments[bmr_number]['comments'].append({
                'type': comment['comment_type'],
                'phase': comment['phase'],
                'date': comment['date'],
                'comments': comment['comments'],
                'status': comment['status']
            })
    except Exception as e:
        return HttpResponse(f"Error retrieving comments data: {e}", content_type="text/plain")
    
    # Create Word document
    doc = Document()
    
    # Add company header and styling
    styles = doc.styles
    
    # Create title style (check if exists first)
    try:
        title_style = styles['KPI Title']
    except KeyError:
        title_style = styles.add_style('KPI Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(18)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    # Create subtitle style (check if exists first)
    try:
        subtitle_style = styles['KPI Subtitle']
    except KeyError:
        subtitle_style = styles.add_style('KPI Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Arial'
        subtitle_font.size = Pt(12)
        subtitle_font.italic = True
    
    # Add company title
    title = doc.add_paragraph('Kampala Pharmaceutical Industries (KPI)', style='KPI Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add report title
    subtitle = doc.add_paragraph('Comments & Observations Report', style='KPI Subtitle')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date and generation info
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    doc.add_paragraph(f'Report prepared by: {request.user.get_full_name()}')
    doc.add_paragraph(f'Total BMRs with comments: {len(bmr_comments)}')
    
    # Add horizontal line
    doc.add_paragraph('_' * 70)
    
    # BMR section style (check if exists first)
    try:
        bmr_style = styles['BMR Heading']
    except KeyError:
        bmr_style = styles.add_style('BMR Heading', WD_STYLE_TYPE.PARAGRAPH)
        bmr_font = bmr_style.font
        bmr_font.name = 'Arial'
        bmr_font.size = Pt(14)
        bmr_font.bold = True
        bmr_font.color.rgb = RGBColor(0, 102, 0)  # Dark green
    
    # Add each BMR section
    for bmr_number, bmr_data in sorted(bmr_comments.items()):
        # Add BMR heading
        doc.add_heading(f"BMR: {bmr_number} - {bmr_data['product']}", level=1).style = bmr_style
        
        # Create table for comments
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set table header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Phase'
        hdr_cells[1].text = 'Comment Type'
        hdr_cells[2].text = 'Date'
        hdr_cells[3].text = 'Comments'
        
        # Make header bold
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add comments to table
        for comment in sorted(bmr_data['comments'], key=lambda x: x['date'] if x['date'] else datetime.min):
            try:
                row_cells = table.add_row().cells
                row_cells[0].text = str(comment['phase']) if comment['phase'] else 'N/A'
                row_cells[1].text = str(comment['type']) if comment['type'] else 'N/A'
                row_cells[2].text = comment['date'].strftime('%Y-%m-%d %H:%M') if comment['date'] else 'N/A'
                row_cells[3].text = str(comment['comments']) if comment['comments'] else 'No comment'
            except Exception as e:
                # Skip problematic rows but continue processing
                continue
        
        # Add spacing after table
        doc.add_paragraph('')
    
    # Add footer
    doc.add_paragraph('_' * 70)
    footer = doc.add_paragraph('This is an automatically generated report from the KPI Operations System.')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="KPI_Comments_Report_{timestamp}.docx"'
    
    doc.save(response)
    return response

@login_required
def export_comments_excel(request):
    """Export comments to Excel format with role-based filtering"""
    try:
        import pandas as pd
        from io import BytesIO
    except ImportError:
        return HttpResponse("pandas library not installed. Please install it to use Excel export.", 
                          content_type="text/plain")
    
    # Get filtered comments data
    raw_comments = get_filtered_comments_data(request)
    
    # Format data for Excel
    comments_data = []
    for comment in raw_comments:
        comments_data.append({
            'BMR Number': comment['bmr_number'],
            'Product': comment['product'],
            'Comment Type': comment['comment_type'],
            'Phase': comment['phase'],
            'Date': comment['date'].strftime('%Y-%m-%d %H:%M:%S') if comment['date'] else '',
            'Comments': comment['comments'],
            'Status': comment['status']
        })
    
    # Create Excel file
    output = BytesIO()
    if comments_data:
        df = pd.DataFrame(comments_data)
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Comments Report', index=False)
    else:
        # Create empty dataframe with headers
        df = pd.DataFrame(columns=['BMR Number', 'Product', 'Comment Type', 'Phase', 'User', 'User Role', 'Date', 'Comments', 'Status'])
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Comments Report', index=False)
    
    output.seek(0)
    
    # Prepare response
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="KPI_Comments_Report_{timestamp}.xlsx"'
    
    return response

@login_required
def bmr_comments_detail(request, bmr_id):
    """Detailed view of all comments for a specific BMR with role-based access"""
    bmr = get_object_or_404(BMR, id=bmr_id)
    
    # Check if user has access to this BMR
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    if not is_admin:
        # Operators can view BMRs in production states, others need involvement
        if bmr.status in ['approved', 'in_production', 'completed']:
            # Allow access to BMRs in production states for all authenticated users
            pass
        else:
            # For non-production BMRs, check user involvement
            user_bmrs = BMR.objects.filter(
                Q(created_by=request.user) | 
                Q(approved_by=request.user)
            )
            
            # Also check if user was involved in any phases of this BMR
            user_phases = BatchPhaseExecution.objects.filter(
                Q(started_by=request.user) | Q(completed_by=request.user),
                bmr=bmr
            )
            
            # Also check signatures
            user_signatures = BMRSignature.objects.filter(
                bmr=bmr,
                signed_by=request.user
            )
            
            # If user has no connection to this BMR, deny access
            if not (bmr in user_bmrs or user_phases.exists() or user_signatures.exists()):
                from django.contrib import messages
                messages.error(request, 'Access denied. You can only view BMRs you were involved in.')
                return redirect('reports:comments_report')
    
    # Collect all comments for this BMR
    comments = []
    
    # BMR level comments
    if bmr.qa_comments:
        comments.append({
            'type': 'BMR QA Comments',
            'phase': 'BMR Creation',
            'user': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
            'role': bmr.created_by.role if bmr.created_by else 'Unknown',
            'date': bmr.created_date,
            'content': bmr.qa_comments,
            'status': bmr.status
        })
    
    if bmr.regulatory_comments:
        comments.append({
            'type': 'BMR Regulatory Comments',
            'phase': 'Regulatory Approval',
            'user': bmr.approved_by.get_full_name() if bmr.approved_by else 'Unknown',
            'role': bmr.approved_by.role if bmr.approved_by else 'Unknown',
            'date': bmr.approved_date or bmr.modified_date,
            'content': bmr.regulatory_comments,
            'status': bmr.status
        })
    
    # Phase comments
    phases = BatchPhaseExecution.objects.filter(bmr=bmr).select_related('phase', 'started_by', 'completed_by')
    
    for phase in phases:
        if phase.operator_comments:
            comments.append({
                'type': 'Operator Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 
                        phase.started_by.get_full_name() if phase.started_by else 'Unknown',
                'role': (phase.completed_by.role if phase.completed_by else 
                        phase.started_by.role if phase.started_by else 'Unknown'),
                'date': phase.completed_date or phase.started_date or phase.created_date,
                'content': phase.operator_comments,
                'status': phase.status
            })
        
        if phase.qa_comments:
            comments.append({
                'type': 'Phase QA Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'content': phase.qa_comments,
                'status': phase.status
            })
        
        if phase.rejection_reason:
            comments.append({
                'type': 'Rejection Reason',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'content': phase.rejection_reason,
                'status': phase.status
            })
    
    # Signature comments
    signatures = BMRSignature.objects.filter(bmr=bmr).select_related('signed_by')
    
    for signature in signatures:
        if signature.comments:
            comments.append({
                'type': 'Electronic Signature',
                'phase': f"Signature - {signature.signed_by_role}",
                'user': signature.signed_by.get_full_name() if signature.signed_by else 'Unknown',
                'role': signature.signed_by_role,
                'date': signature.signed_date,
                'content': signature.comments,
                'status': 'Signed'
            })
    
    # Sort by date
    comments.sort(key=lambda x: x['date'] or datetime.min)
    
    context = {
        'bmr': bmr,
        'comments': comments,
        'total_comments': len(comments),
        'is_admin': is_admin,
        'user_role': request.user.role
    }
    
    return render(request, 'reports/bmr_comments_detail.html', context)
